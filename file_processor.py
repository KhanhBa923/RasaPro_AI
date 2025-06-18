import os
from pathlib import Path
import PyPDF2
from docx import Document
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import AzureOpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv


load_dotenv()


def extract_text_from_word(docx_path):
    doc = Document(docx_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text

def extract_text_from_excel(excel_path):
    excel_file = pd.ExcelFile(excel_path)
    text = ""
    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(excel_path, sheet_name=sheet_name)
        text += f"\n--- Sheet: {sheet_name} ---\n"
        text += " | ".join(df.columns.astype(str)) + "\n"
        for _, row in df.iterrows():
            text += " | ".join(row.astype(str)) + "\n"
    return text

def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            if page.extract_text():
                text += page.extract_text()
    return text

def process_documents_to_vector_store():
    docs_folder = Path("./docs")
    all_documents = []

    file_processors = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_word,
        '.doc': extract_text_from_word,
        '.xlsx': extract_text_from_excel,
        '.xls': extract_text_from_excel
    }

    if not docs_folder.exists():
        print("📁 Thư mục 'docs' không tồn tại.")
        return
    try:
        for file_path in docs_folder.iterdir():
            if file_path.suffix.lower() in file_processors:
                try:
                    processor = file_processors[file_path.suffix.lower()]
                    text = processor(file_path)
                    print(f"✅ Processing: {file_path}")
                    if not text.strip():
                        print(f"⚠️ File {file_path} không có nội dung.")
                        continue

                    document = {
                        'content': text,
                        'source': str(file_path),
                        'file_type': file_path.suffix.lower()
                    }
                    all_documents.append(document)

                    print(f"✅ Processed: {file_path}")
                except Exception as e:
                    print(f"❌ Error processing {file_path}: {e}")
    except Exception as e:
        print(f"Error: +{e}")
    if not all_documents:
        print("⚠️ Không có tài liệu hợp lệ để xử lý.")
        return

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = []
    metadatas = []

    for doc in all_documents:
        doc_chunks = text_splitter.split_text(doc['content'])
        chunks.extend(doc_chunks)
        metadatas.extend([{
            'source': doc['source'],
            'file_type': doc['file_type']
        }] * len(doc_chunks))

    if not chunks:
        print("⚠️ Không có đoạn văn bản nào được tạo.")
        return

    embeddings = AzureOpenAIEmbeddings(
    azure_deployment=os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION")
    )
    vector_store = FAISS.from_texts(chunks, embeddings, metadatas=metadatas)
    vector_store.save_local("./vector_store")

    print(f"✅ Tạo xong vector store với {len(chunks)} đoạn.")

if __name__ == "__main__":
    process_documents_to_vector_store()
